"""Word 登记表写入模块

职责：
- 按模板格式生成每条评论句的 Word 登记表
- 评论句中标志词、作者姓氏、年份自动加粗
"""

import logging
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

from core.llm_analyzer import CommentRecord
from core.pdf_parser import PaperMetadata

logger = logging.getLogger(__name__)


def _extract_surname(full_name: str) -> str:
    """从全名中提取姓氏

    规则：
    - 中文作者：返回全名（中文姓名整体视为姓氏）
    - 英文作者：提取 last name（姓）
      - "Smith, John" → "Smith"
      - "John Smith" → "Smith"
      - "van der Berg, Jan" → "van der Berg"
    """
    if not full_name or not full_name.strip():
        return full_name

    name = full_name.strip()

    # 中文作者：含中文字符则返回全名
    if any('\u4e00' <= c <= '\u9fff' for c in name):
        return name

    # 英文作者：逗号格式 "Last, First" → 取逗号前
    if ',' in name:
        return name.split(',')[0].strip()

    # 英文作者：空格格式 "First Last" 或 "First Middle Last" → 取最后一个词
    # 但需处理 "van der Berg" 等复合姓氏的情况
    parts = name.split()
    if len(parts) == 1:
        return name

    # 常见姓氏前缀
    prefixes = {'van', 'von', 'de', 'del', 'della', 'der', 'den', 'di', 'la', 'le', 'el', 'al', 'bin', 'ibn'}

    # 从后往前找姓氏起始位置
    # 假设第一个词是名字，从第二个词开始看是否包含前缀
    surname_start = len(parts) - 1
    for i in range(1, len(parts) - 1):
        if parts[i].lower() in prefixes:
            surname_start = i
            break

    return " ".join(parts[surname_start:])


def _clean_marker(marker: str) -> str:
    """清洗标志词，去除前后多余的介词/助词/be动词

    要求 d：严格按标志词表中的词加粗，不加粗前后带着的介词或别的多余的词。
    例如 LLM 返回 "was first proposed by" → 清洗为 "first proposed"
    """
    if not marker or not marker.strip():
        return marker

    cleaned = marker.strip()

    # 前缀：常见的 be 动词、助动词、介词等（不属于标志词本身）
    prefix_words = [
        'was', 'were', 'is', 'are', 'been', 'being', 'have', 'has', 'had',
        'the', 'a', 'an', 'by', 'in', 'on', 'at', 'to', 'for', 'with',
        'of', 'from', 'that', 'which', 'who', 'it',
    ]

    # 后缀：常见的介词、连词等
    suffix_words = [
        'by', 'in', 'on', 'at', 'to', 'for', 'with', 'of', 'from',
        'that', 'the', 'a', 'an', 'as',
    ]

    # 反复去除前缀
    changed = True
    while changed:
        changed = False
        for pw in prefix_words:
            pattern = re.compile(r'^' + re.escape(pw) + r'\s+', re.IGNORECASE)
            new = pattern.sub('', cleaned)
            if new != cleaned:
                cleaned = new
                changed = True

    # 反复去除后缀
    changed = True
    while changed:
        changed = False
        for sw in suffix_words:
            pattern = re.compile(r'\s+' + re.escape(sw) + r'$', re.IGNORECASE)
            new = pattern.sub('', cleaned)
            if new != cleaned:
                cleaned = new
                changed = True

    return cleaned.strip()


def _add_run(paragraph, text: str, bold: bool = False,
             font_name: str = "Times New Roman", font_size: float = 12,
             east_asia_font: str = "宋体"):
    """添加一个格式化的 run 到段落（减少重复代码）

    默认字体：西文 Times New Roman / 中文 宋体 / 12pt（匹配模板）
    """
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_font)
    return run


def _set_cell_font(cell, font_name="Times New Roman", font_size=10.5):
    """设置单元格默认字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _write_bold_sentence(cell, sentence: str, marker: str, author: str, year: str):
    """在单元格中写入评论句，标志词/作者姓氏/年份加粗

    加粗规则（要求 c/d）：
    - 只加粗标志词（严格按标志词表，不含前后介词等多余词）
    - 只加粗第一作者姓氏（不加粗 et al. 和其他作者）
    - 加粗年份
    """
    cell.text = ""  # 清空
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 收集所有需要加粗的位置
    bold_ranges = []

    # 查找标志词位置（严格匹配，不加粗前后多余的词）
    if marker:
        # 清洗标志词：去除前后可能的多余介词/助词
        cleaned_marker = _clean_marker(marker)
        if cleaned_marker:
            for m in re.finditer(re.escape(cleaned_marker), sentence):
                bold_ranges.append((m.start(), m.end()))

    # 查找作者姓氏位置（只加粗姓氏，不加粗 et al. 和其他作者）
    if author:
        surname = _extract_surname(author)
        if surname:
            for m in re.finditer(re.escape(surname), sentence):
                bold_ranges.append((m.start(), m.end()))

    # 查找年份位置（4位数字年份 + 年代词如 1950s）
    if year:
        year_clean = year.strip()
        if len(year_clean) == 4 and year_clean.isdigit():
            for m in re.finditer(re.escape(year_clean), sentence):
                bold_ranges.append((m.start(), m.end()))
        elif year_clean:
            # 年代词如 "the 1950s"，只加粗数字部分如 "1950s"
            digit_match = re.search(r'\d{4}s?', year_clean)
            if digit_match:
                year_core = digit_match.group()
                for m in re.finditer(re.escape(year_core), sentence):
                    bold_ranges.append((m.start(), m.end()))

    # 合并重叠区间
    bold_ranges.sort()
    merged = []
    for start, end in bold_ranges:
        if merged and start <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))

    # 按区间切分文本，分别添加 run
    pos = 0
    for start, end in merged:
        # 非加粗部分
        if pos < start:
            _add_run(paragraph, sentence[pos:start])
        # 加粗部分
        _add_run(paragraph, sentence[start:end], bold=True)
        pos = end

    # 剩余部分
    if pos < len(sentence):
        _add_run(paragraph, sentence[pos:])


def _write_cell(cell, text: str, bold=False, font_size: float = 12):
    """写入单元格文本（默认 Times New Roman/宋体 12pt）"""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _add_run(paragraph, str(text), bold=bold, font_size=font_size)


def _format_ref_for_word(record: CommentRecord) -> str:
    """格式化被评文献引用（Word 登记表格式）"""
    ep = record.被评文献
    authors = ", ".join(ep.全部作者列表) if ep.全部作者列表 else ep.第一作者
    parts = [authors]
    if ep.文章名:
        parts[0] += f". {ep.文章名}[J]."
    if ep.期刊名称:
        parts.append(f" {ep.期刊名称},")
    if ep.年份:
        parts.append(f" {ep.年份},")
    vol = ""
    if ep.卷:
        vol = ep.卷
        if ep.期:
            vol += f"({ep.期})"
        vol += ":"
    if vol:
        parts.append(f" {vol}")
    if ep.起止页码:
        parts.append(ep.起止页码)
    return "".join(parts) + "."


def _set_cell_valign(cell, align: str = "center"):
    """设置单元格垂直对齐"""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = tc.makeelement(qn('w:tcPr'), {})
        tc.insert(0, tcPr)
    vAlign = tcPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = tcPr.makeelement(qn('w:vAlign'), {})
        tcPr.append(vAlign)
    vAlign.set(qn('w:val'), align)


def _set_row_height(row, height_twips: int, rule: str = "atLeast"):
    """设置行高（单位 twips）"""
    trPr = row._tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = row._tr.makeelement(qn('w:trPr'), {})
        row._tr.insert(0, trPr)
    trHeight = trPr.find(qn('w:trHeight'))
    if trHeight is None:
        trHeight = trPr.makeelement(qn('w:trHeight'), {})
        trPr.append(trHeight)
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), rule)


def write_word(
    output_path: str,
    record: CommentRecord,
    metadata: PaperMetadata,
    index: int = 1,
    institution_info: dict | None = None,
    provider: str = "",
) -> str:
    """生成单条评论句的 Word 登记表（严格匹配模板格式）

    模板来源：【输出2】学术评论句登记表标准格式.docx
    - 页面：A4 纵向，四边距 1.27cm
    - 字体：西文 Times New Roman / 中文 宋体 / 12pt（标题 14pt）
    - 列宽：2.52 / 4.37 / 8.81 / 1.24 / 1.37 cm（共5列）
    - 行高：按模板逐行指定，rule=atLeast
    - 所有单元格垂直居中
    """
    doc = Document()

    # ── 页面设置：A4 纵向，四边距 1.27cm ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    # 设置默认字体样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 创建表格 (16行 x 5列) ──
    table = doc.add_table(rows=16, cols=5, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格宽度 10377 dxa（匹配模板）
    tblPr = table._tbl.find(qn('w:tblPr'))
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = tblPr.makeelement(qn('w:tblW'), {})
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '10377')
    tblW.set(qn('w:type'), 'dxa')

    # ── 设置列宽（匹配模板：1427/2476/4994/702/778 twips）──
    COL_WIDTHS = [1427, 2476, 4994, 702, 778]
    tblGrid = table._tbl.find(qn('w:tblGrid'))
    for i, col in enumerate(tblGrid.findall(qn('w:gridCol'))):
        col.set(qn('w:w'), str(COL_WIDTHS[i]))

    # ── 设置行高（匹配模板，单位 twips）──
    ROW_HEIGHTS = [580, 1012, 976, 648, 688, 480, 615, 442, 472, 1609, 2541, 648, 648, 541, 494, 494]
    for i, row in enumerate(table.rows):
        _set_row_height(row, ROW_HEIGHTS[i])

    # ── 所有单元格垂直居中 ──
    for row in table.rows:
        for cell in row.cells:
            _set_cell_valign(cell, "center")

    ep = record.被评文献
    inst = institution_info or {}

    eval_institution = ep.第一作者机构 or inst.get("institution", "")
    eval_country = ep.第一作者国家 or inst.get("country", "")

    # ── 行0: 标题行 ──
    # 合并列0-1（2列）
    cell_a0 = table.cell(0, 0)
    cell_b0 = table.cell(0, 1)
    cell_a0.merge(cell_b0)
    _write_cell(cell_a0, "2026年学术评论句登记表", bold=True, font_size=14)

    # 合并列2-4（3列）
    cell_c0 = table.cell(0, 2)
    cell_e0 = table.cell(0, 4)
    cell_c0.merge(cell_e0)
    _write_cell(cell_c0, "               流水号：      /         ")

    # ── 定义行内容 ──
    authors_str = metadata.authors_str

    # 期刊格式：GB/T 7714
    journal = metadata.journal_cn or metadata.journal_en
    vol_info = ""
    if metadata.volume:
        vol_info = f"{metadata.volume}"
        if metadata.issue:
            vol_info += f"({metadata.issue})"
    pages = metadata.pages
    journal_str = f"{journal}, {metadata.year}, {vol_info}: {pages}." if journal else ""

    row_definitions = [
        # (行号, 标题, 内容)
        (1, "施评文献作者", authors_str),
        (2, "期刊名称，年，卷，期，页", journal_str),
        (3, "机构中文名称", metadata.institution_cn),
        (4, "机构英文名称", metadata.institution_en or ""),
        (5, "中文题目", metadata.title_cn),
        (6, "英文题目", metadata.title_en),
        (7, "论文关键词", "/"),
        (8, "题目关键词", "/"),
        (9, "学术评论句", None),  # 特殊处理
        (10, "被评文献", _format_ref_for_word(record)),
        (11, "其他被评文献", ""),
        (12, "被评作者机构", f"{eval_institution}, {eval_country}" if eval_institution else ""),
        (13, "标志词", record.标志词),
        (14, "提供者", provider),
    ]

    for row_idx, title, content in row_definitions:
        # A列（列0）：标题
        _write_cell(table.cell(row_idx, 0), title)

        # B-C列（列1-2）合并：内容
        cell_b = table.cell(row_idx, 1)
        cell_c = table.cell(row_idx, 2)
        cell_b.merge(cell_c)

        if row_idx == 9:
            # 学术评论句：标志词/作者姓氏/年份加粗（要求 c/d）
            _write_bold_sentence(
                cell_b,
                record.评论句原文,
                record.标志词,
                ep.第一作者,
                ep.年份,
            )
        elif row_idx == 2:
            # 期刊栏：年份加粗（要求 b）
            _write_bold_journal(cell_b, journal_str, metadata.year)
        elif row_idx == 10:
            # 被评文献栏：第一作者姓氏 + 年份加粗（要求 e）
            _write_bold_ref(cell_b, content or "", ep.第一作者, ep.年份)
        else:
            _write_cell(cell_b, content or "")

        # D列（列3）
        if row_idx == 1:
            _write_cell(table.cell(row_idx, 3), "第1次反馈")
        elif row_idx in (9, 10, 11):
            _write_cell(table.cell(row_idx, 3), "/")
        else:
            _write_cell(table.cell(row_idx, 3), "/")

        # E列（列4）
        if row_idx == 1:
            _write_cell(table.cell(row_idx, 4), "第2次反馈")
        elif row_idx in (9, 10, 11):
            _write_cell(table.cell(row_idx, 4), "/")
        else:
            _write_cell(table.cell(row_idx, 4), "/")

    # ── 行15: 备注行（5列合并）──
    cell_all = table.cell(15, 0)
    cell_end = table.cell(15, 4)
    cell_all.merge(cell_end)
    cell_all.text = ""
    p = cell_all.paragraphs[0]
    _add_run(p, "其他：你这个评论句有（")
    _add_run(p, "/")
    _add_run(p, "）个被评文献，请提供另（")
    _add_run(p, "/")
    _add_run(p, "）张登记表。")
    _add_run(p, "【注意重新提交修正后的Excel表、登记表及原文】")

    # ── 在行1中，给第一作者姓氏加粗（要求 a）──
    if metadata.first_author:
        cell_author = table.cell(1, 1)
        _write_bold_author(cell_author, authors_str, metadata.first_author)

    doc.save(output_path)
    logger.info(f"Word 登记表已保存: {output_path}")
    return output_path


def _write_bold_author(cell, authors_str: str, first_author: str):
    """在作者列表中将第一作者姓氏加粗（要求 a）

    只加粗第一作者的姓氏，不加粗名字。
    """
    cell.text = ""
    paragraph = cell.paragraphs[0]

    surname = _extract_surname(first_author)
    if not surname or surname not in authors_str:
        # 姓氏无法提取或不在字符串中，不加粗
        _add_run(paragraph, authors_str)
        return

    # 找到姓氏在作者字符串中的位置（只加粗第一次出现）
    idx = authors_str.index(surname)
    # 前缀
    if idx > 0:
        _add_run(paragraph, authors_str[:idx])
    # 姓氏加粗
    _add_run(paragraph, surname, bold=True)
    # 后缀
    end_idx = idx + len(surname)
    if end_idx < len(authors_str):
        _add_run(paragraph, authors_str[end_idx:])


def _write_bold_journal(cell, journal_str: str, year: str):
    """在期刊信息栏中将年份加粗（要求 b）

    例如 "Journal Name, 2024, 15(3): 100-110." → 2024 加粗
    """
    cell.text = ""
    paragraph = cell.paragraphs[0]

    if not year or not journal_str:
        _add_run(paragraph, journal_str or "")
        return

    year_clean = year.strip()
    if year_clean not in journal_str:
        _add_run(paragraph, journal_str)
        return

    idx = journal_str.index(year_clean)
    # 年份前的部分
    if idx > 0:
        _add_run(paragraph, journal_str[:idx])
    # 年份加粗
    _add_run(paragraph, year_clean, bold=True)
    # 年份后的部分
    end_idx = idx + len(year_clean)
    if end_idx < len(journal_str):
        _add_run(paragraph, journal_str[end_idx:])


def _write_bold_ref(cell, ref_str: str, first_author: str, year: str):
    """在被评文献栏中将第一作者姓氏和年份加粗（要求 e）

    只加粗第一作者的姓氏和年份，不加粗名字，非第一作者不加粗。
    """
    cell.text = ""
    paragraph = cell.paragraphs[0]

    if not ref_str:
        return

    # 收集需要加粗的区间
    bold_ranges = []

    # 查找第一作者姓氏
    surname = _extract_surname(first_author) if first_author else ""
    if surname and surname in ref_str:
        # 只加粗第一次出现（即作者列表开头的那个）
        idx = ref_str.index(surname)
        bold_ranges.append((idx, idx + len(surname)))

    # 查找年份（4位数字）
    if year:
        year_clean = year.strip()
        for m in re.finditer(re.escape(year_clean), ref_str):
            bold_ranges.append((m.start(), m.end()))

    if not bold_ranges:
        _add_run(paragraph, ref_str)
        return

    # 合并重叠区间
    bold_ranges.sort()
    merged = []
    for start, end in bold_ranges:
        if merged and start <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))

    # 按区间切分输出
    pos = 0
    for start, end in merged:
        if pos < start:
            _add_run(paragraph, ref_str[pos:start])
        _add_run(paragraph, ref_str[start:end], bold=True)
        pos = end
    if pos < len(ref_str):
        _add_run(paragraph, ref_str[pos:])
